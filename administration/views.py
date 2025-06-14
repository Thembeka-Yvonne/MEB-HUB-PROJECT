from django.shortcuts import render,redirect
from login.models import Campus, Admin, Student, RegisteredStudent
from bus.models import Route,Bus_schedule,Bus,Bus_Stats
from .models import Admin_Action
from events.models import RSVP
from django.urls import reverse
from django.contrib import messages
from django.http import HttpResponseRedirect,HttpResponse
from datetime import datetime,timedelta
from datetime import date
from django.forms import Form
from events.models import Event, ArchivedEvents
from django.utils import timezone
from django.db.models.functions import TruncDate
from django.db.models import Count
from django.utils.timezone import now, timedelta
from reportlab.pdfgen import canvas
from docx import Document
import csv
from io import BytesIO
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from django.db.models import Sum
import json
from django.db.models.functions import TruncDay, TruncWeek, TruncMonth
from django.utils.safestring import mark_safe
from django.utils.timezone import localtime

# Create your views here.
def admin_home(request):
 if 'admin_id' in request.session:
  admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
  today = timezone.localdate()
  actions = Admin_Action.objects.annotate(action_date=TruncDate('datetime')).filter(admin_id=admin.admin_id,action_date=today)
  events = Event.objects.filter(date__year=today.year, date__month=today.month)
  cnt_routes=Route.objects.count()

  icon = request.GET.get('type')
  action_type = 'all'
  if icon == 'bi bi-map':
      action_type = 'Map Actions'
  elif icon == 'bi bi-calendar':
      action_type = 'Event Actions'
  elif icon == 'bi bi-bus-front':
      action_type = 'Bus Actions'
  else :
      action_type="All Actions"

  context={
    "admin":admin,
    "actions" :actions,
    "events" :events,
    "initials": f"{admin.name[0].upper()}{admin.surname[0].upper()}",
    "cnt_routes":cnt_routes,
    "action_type":action_type
  }
  return render(request,"admin/admin.html",context
  )
 else:
   return redirect('account:login')

#campus functionalilities

#bus functionalities
def bus_menu(request):
  initials=request.session.get('initials')
  return render(request,"admin/buses/bus_menu.html",{'initials':initials})

def add_all_campuses_view(request):
  list = Route.objects.all()
  initials = request.session.get("initials")
  return render(request,"admin/buses/add_all_camp.html",{
    "list": list,"initials":initials
  })
  
def add_bus_schedule(request,code):
  
  schedule_c = Route.objects.all().get(schedule_code=code)
  bus_list = Bus.objects.all()
  initials = request.session.get("initials")
  
  errors = {}
  
  if request.method == 'POST':

    bus = Bus.objects.all().get(bus_id=int(request.POST['bus']))
    start_time = request.POST['start_time']
    last_time = request.POST['last_time']
    duration = int(request.POST['duration'])
    
    # Basic field validation
    if not bus or not start_time or not last_time or not duration:
        errors['form'] = "All fields are required."
    
    s_time = datetime.strptime(start_time, '%H:%M')
    l_time = datetime.strptime(last_time, '%H:%M')
    
    if l_time < s_time:
      errors['time error'] = "Last time should not be earlier than start time"
      
    # if duration > 2 or duration < 1 :
    #   errors['duration error'] = "Duration should be between 1 and 2"
    
    n_time = datetime.strptime(start_time, '%H:%M')
      
    if not errors:
      Bus_schedule.objects.filter(schedule_code=code).delete()
      n_time = datetime.strptime(start_time, '%H:%M')
      while n_time < l_time:
        n_time = s_time + timedelta(minutes=duration)
        
        bus_s = Bus_schedule(departure=schedule_c.campus1,destination=schedule_c.campus2,
                               departure_time=s_time.time(),arrival_time=n_time.time(),duration=duration,
                               bus_id=bus,schedule_code=schedule_c)
          
        bus_s.save()
          
        s_time = n_time + timedelta(minutes=duration)
        n_time = s_time + timedelta(minutes=duration)
        
        if n_time > l_time:
          break
    
        bus_s = Bus_schedule(departure=schedule_c.campus2,destination=schedule_c.campus1,
                          departure_time=s_time.time(),arrival_time=n_time.time(),duration=duration,
                          bus_id=bus,schedule_code=schedule_c)
        bus_s.save()
          
        s_time = n_time + timedelta(minutes=duration)
          
      admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
      addAction(admin_id=admin,record_type="Generated bus schedule",icon="bi bi-bus-front")
      return HttpResponseRedirect(redirect_to='/administration/bus_menu')
    else:
        return render(request,"admin/buses/add_schedule.html",{
          "schedule": schedule_c,
          "bus_list": bus_list,"initials":initials,
          "errors": errors
        })
  
  return render(request,"admin/buses/add_schedule.html",{
        "schedule": schedule_c,
        "bus_list": bus_list,"initials":initials,
        "errors": errors
      })
  
def events_menu(request):
  expired_events = Event.objects.filter(date__lt=date.today()) #automatically delete events
  for event in expired_events:
        event.delete()

  admin_id = request.session.get('admin_id')
  initials = request.session.get("initials")
  total_upcoming_events = Event.objects.count()
  total_past_events = ArchivedEvents.objects.count()

  filtered_events=None
  start_date=request.GET.get('start_date')
  end_date=request.GET.get('end_date')
  if start_date and end_date:
      filtered_events=Event.objects.filter(date__range=[start_date,end_date])




  filtered_events_admin = Event.objects.filter(admin_id_id=admin_id)

  if admin_id is None:
    return redirect('admin_home')  # or handle expired session

  admin = Admin.objects.select_related('campus_id').get(admin_id=admin_id)
  return render(request,"admin/events/events_menu.html",{'admin':admin,'initials':initials,'total_upcoming_events': total_upcoming_events,
        'total_past_events': total_past_events,
        'filtered_events': filtered_events,'filtered_events_admin':filtered_events_admin})

def download_filtered_events(request):
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    events = Event.objects.filter(date__range=[start_date, end_date])

    # Create PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{start_date}_to_{end_date}_events.pdf"'

    # Create the PDF object
    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4
    y = height - inch

    # Title
    p.setFont("Helvetica-Bold", 16)
    p.drawCentredString(width / 2.0, y, f"Filtered Events from {start_date} to {end_date}")
    y -= 0.5 * inch

    # Table Header
    p.setFont("Helvetica-Bold", 12)
    p.drawString(inch, y, "Event Title")
    p.drawString(3 * inch, y, "Date")
    p.drawString(4.5 * inch, y, "Location")
    y -= 0.3 * inch

    # Table Rows
    p.setFont("Helvetica", 11)
    for event in events:
      if y < inch:  # Start a new page if too low
        p.showPage()
        y = height - inch
      p.drawString(inch, y, event.title)
      p.drawString(3 * inch, y, str(event.date))
      p.drawString(4.5 * inch, y, event.location)
      y -= 0.25 * inch

      # No events message
      if not events:
        p.drawString(inch, y, "No events found for the selected date range.")

    p.showPage()
    p.save()

    return response


def download_filtered_events_csv(request):
  start_date = request.GET.get('start_date')
  end_date = request.GET.get('end_date')
  events = Event.objects.filter(date__range=[start_date, end_date])

  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = 'attachment; filename="Filtered Events from {start_date} to {end_date}.csv"'
  writer = csv.writer(response)
  writer.writerow(['Title', 'Date', 'Location'])
  for event in events:
    writer.writerow([event.title, event.date, event.location])
  return response

def download_filtered_events_admin(request):
  admin_id = request.session.get('admin_id')
  events = Event.objects.filter(admin_id_id=admin_id)
  admin=Admin.objects.get(admin_id=admin_id)

  # Create PDF response
  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = f'attachment; filename="events_by_{admin_id}.pdf"'

  # Create the PDF object
  p = canvas.Canvas(response, pagesize=A4)
  width, height = A4
  y = height - inch

  # Title
  p.setFont("Helvetica-Bold", 16)
  p.drawCentredString(width / 2.0, y, f" Events Posted by {admin.name} {admin.surname}({admin.admin_id})")
  y -= 0.5 * inch

  # Table Header
  p.setFont("Helvetica-Bold", 12)
  p.drawString(inch, y, "Event Title")
  p.drawString(3 * inch, y, "Date")
  p.drawString(4.5 * inch, y, "Location")
  y -= 0.3 * inch

  # Table Rows
  p.setFont("Helvetica", 11)
  for event in events:
    if y < inch:  # Start a new page if too low
      p.showPage()
      y = height - inch
    p.drawString(inch, y, event.title)
    p.drawString(3 * inch, y, str(event.date))
    p.drawString(4.5 * inch, y, event.location)
    y -= 0.25 * inch

    # No events message
    if not events:
      p.drawString(inch, y, "No events found for the selected date range.")

  p.showPage()
  p.save()

  return response

def addAction(admin_id: int,record_type: str,icon: str):
  action = Admin_Action(action_type=record_type,admin_id=admin_id, icon=icon,
                        datetime = timezone.now())
  action.save()
  

def view_all_actions(request):
  admin = Admin.objects.all().get(admin_id=request.session['admin_id'])
  today = timezone.now().date()
  actions = Admin_Action.objects.annotate(action_date=TruncDate('datetime')).filter(admin_id=admin.admin_id,
                                                                                    action_date=today)
  events = Event.objects.filter(date__year=today.year, date__month=today.month)
  cnt_routes = Route.objects.count()
  context = {
    "admin": admin,
    "actions": actions,
    "events": events,
    "initials": f"{admin.name[0].upper()}{admin.surname[0].upper()}",
    "cnt_routes": cnt_routes
  }
  
  return render(request,"admin/admin.html",context)

def add_bus(request):
  if request.method == 'POST':
    form = Form(request.POST)
    if form.is_valid():
      bus_name = request.POST["bus_name"]
      campus_id = request.POST["campus"]

      campus = Campus.objects.all().get(campus_id=campus_id)
      bus = Bus(bus_name=bus_name,campus_id=campus)

      bus.save()
      
      admin = Admin.objects.all().get(admin_id=request.session['admin_id'])


      
      addAction(admin_id=admin,record_type="Added a new bus",icon="bi bi-bus-front")
      
      return redirect('view_all_actions')
    
    else:
       return  render(request,"admin/buses/add_bus_html",{
         "form": form
       })

  campus_list = Campus.objects.all()
  initials = request.session.get("initials")
  return render(request,"admin/buses/add_bus.html",{
    "campus_list": campus_list,"initials":initials
  })


def remove_bus(request):
   if request.method == 'POST':
      form = Form(request.POST)
      initials = request.session.get("initials")

      if form.is_valid:
        bus_id = request.POST["bus"]

        bus = Bus.objects.all().get(bus_id=bus_id)
        
        bus.delete()
        
        admin = Admin.objects.all().get(admin_id=request.session['admin_id'])


        addAction(admin_id=admin,record_type="Bus has been removed",icon="bi bi-x-circle")
        
        return redirect('view_all_actions')
      else:
        return render(request,"admin/buses/remove_bus.html",{
          "form": form,"initials":initials
        })
   initials = request.session.get("initials")

   bus_list = Bus.objects.all()
   return render(request,"admin/buses/remove_bus.html",{
     "bus_list": bus_list,"initials":initials
   })

def bus_schedule_stats(request):
  total_views = Bus_Stats.objects.count()

  # Views per schedule
  views_per_schedule = Bus_Stats.objects.values(
    'schedule_code__schedule_code',
    'schedule_code__campus1',
    'schedule_code__campus2'
  ).annotate(view_count=Count('id')).order_by('-view_count')

  # Views by day
  views_by_day = Bus_Stats.objects.annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Views in the last 7 days
  seven_days_ago = timezone.now() - timedelta(days=7)
  recent_views = Bus_Stats.objects.filter(viewed_at__gte=seven_days_ago).annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Prepare data for charts
  schedule_labels = [
    f"{item['schedule_code__schedule_code']} - {item['schedule_code__campus1']} and {item['schedule_code__campus2']}"
    for item in views_per_schedule
  ]
  schedule_counts = [item['view_count'] for item in views_per_schedule]

  day_labels = [item['date'] for item in views_by_day]
  day_counts = [item['count'] for item in views_by_day]

  recent_labels = [item['date'] for item in recent_views]
  recent_counts = [item['count'] for item in recent_views]

  context = {
    'total_views': total_views,
    'views_per_schedule': list(zip(schedule_labels, schedule_counts)),
    'views_by_day': list(zip(day_labels, day_counts)),
    'recent_views': list(zip(recent_labels, recent_counts)),
    'most_viewed': views_per_schedule[0] if views_per_schedule else None,
  }

  return render(request, 'admin/buses/schedule_stats.html', context)

def user_management(request):
    initials=request.session.get("initials")
    return render(request,'admin/user_management.html',{"initials":initials})

def analytics(request):
    initials=request.session.get("initials")
    return render(request,'admin/analytics.html',{'initials':initials})
      
  
def bus_schedule_stats(request):
  total_views = Bus_Stats.objects.count()

  # Views per schedule
  views_per_schedule = Bus_Stats.objects.values(
    'schedule_code__schedule_code',
    'schedule_code__campus1',
    'schedule_code__campus2'
  ).annotate(view_count=Count('id')).order_by('-view_count')

  # Views by day
  views_by_day = Bus_Stats.objects.annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Views in the last 7 days
  seven_days_ago = timezone.now() - timedelta(days=7)
  recent_views = Bus_Stats.objects.filter(viewed_at__gte=seven_days_ago).annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Prepare data for charts
  schedule_labels = [
    f"{item['schedule_code__schedule_code']} - {item['schedule_code__campus1']} and {item['schedule_code__campus2']}"
    for item in views_per_schedule
  ]
  schedule_counts = [item['view_count'] for item in views_per_schedule]

  day_labels = [item['date'] for item in views_by_day]
  day_counts = [item['count'] for item in views_by_day]

  recent_labels = [item['date'] for item in recent_views]
  recent_counts = [item['count'] for item in recent_views]



  context = {
    'total_views': total_views,
    'views_per_schedule': list(zip(schedule_labels, schedule_counts)),
    'views_by_day': list(zip(day_labels, day_counts)),
    'recent_views': list(zip(recent_labels, recent_counts)),
    'most_viewed': views_per_schedule[0] if views_per_schedule else None,
  }

  return render(request, 'admin/buses/schedule_stats.html', context)

def user_management(request):
    initials=request.session.get("initials")
    students = Student.objects.all()
    total_students = students.count()
    campus_counts = students.values('campus_id__campus_name').annotate(count=Count('studentNumber')).order_by('-count')
    recent_students = Student.objects.order_by('-login_time')[:5]

    student = None
    studentEmail = request.GET.get('studentEmail')  # from URL or form

    campuses = Campus.objects.all()
    campus_id=request.GET.get('campus_id')
    if campus_id:
        registered_students = RegisteredStudent.objects.filter(campus_id=campus_id)
        student_numbers = registered_students.values_list('studentNumber', flat=True)
        students = students.filter(studentNumber__in=student_numbers)



    if studentEmail:
      try:
        student = Student.objects.get(studentEmail=studentEmail)
      except Student.DoesNotExist:
        student = None

    context={
      'total_students':total_students,
      'initials':initials,
      'campus_counts':campus_counts,
      'recent_students':recent_students,
      'student':student,
      'campuses': campuses,
      'students':students
    }
    return render(request,'admin/user_management.html',context)

def analytics(request):
    initials=request.session.get("initials")

    # Get initials from session
    initials = request.session.get("initials")

    # Get both stats
    bus_stats = bus_schedule_stats()
    event_stats = events_stats()
    user_activity_data = get_user_activity_data()

    context = {
        **bus_stats,
        **event_stats,
        **user_activity_data,
        "initials": initials
    }

    return render(request,'admin/analytics.html',context)
      
  
def bus_schedule_stats():
  total_views = Bus_Stats.objects.count()

  # Views per schedule
  views_per_schedule = Bus_Stats.objects.values(
    'schedule_code__schedule_code',
    'schedule_code__campus1',
    'schedule_code__campus2'
  ).annotate(view_count=Count('id')).order_by('-view_count')

  # Views by day
  views_by_day = Bus_Stats.objects.annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Views in the last 7 days
  seven_days_ago = timezone.now() - timedelta(days=7)
  recent_views = Bus_Stats.objects.filter(viewed_at__gte=seven_days_ago).annotate(
    date=TruncDate('viewed_at')
  ).values('date').annotate(count=Count('id')).order_by('date')

  # Prepare data for charts
  schedule_labels = [
    f"{item['schedule_code__schedule_code']} - {item['schedule_code__campus1']} and {item['schedule_code__campus2']}"
    for item in views_per_schedule
  ]
  schedule_counts = [item['view_count'] for item in views_per_schedule]

  day_labels = [item['date'] for item in views_by_day]
  day_counts = [item['count'] for item in views_by_day]

  recent_labels = [item['date'] for item in recent_views]
  recent_counts = [item['count'] for item in recent_views]



  context = {
    'total_views': total_views,
    'views_per_schedule': list(zip(schedule_labels, schedule_counts)),
    'views_by_day': list(zip(day_labels, day_counts)),
    'recent_views': list(zip(recent_labels, recent_counts)),
    'most_viewed': views_per_schedule[0] if views_per_schedule else None,
  }

  return context


def student_report(request):
    students = Student.objects.all()

    total_students = students.count()
    bus_users = Bus_Stats.objects.count()
    non_bus_users = total_students - bus_users

    campus_counts = students.values('campus_id__campus_name').annotate(count=Count('studentNumber')).order_by('-count')

    latest_login = students.order_by('-login_time').first()

    initials=request.session.get("initials")

    context = {
      'students': students,
      'total_students': total_students,
      'bus_users': bus_users,
      'non_bus_users': non_bus_users,
      'campus_counts': campus_counts,
      'latest_login': latest_login,
      'initials': initials
    }
    return render(request, 'admin/student_report.html', context)


def student_report_pdf(request):
  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = 'attachment; filename="student_report.pdf"'

  p = canvas.Canvas(response)
  p.setFont("Helvetica", 12)

  # Metrics
  total_students = Student.objects.count()
  bus_users = Bus_Stats.objects.count()
  non_bus_users = Bus_Stats.objects.count()
  campus_counts = Student.objects.values('campus_id__campus_name').annotate(count=Count('studentNumber'))

  y = 800
  p.drawString(100, y, "Student Report")
  y -= 30

  p.drawString(50, y, f"Total Students: {total_students}")
  y -= 20
  p.drawString(50, y, f"Students Using Bus: {bus_users}")
  y -= 20
  p.drawString(50, y, f"Students Not Using Bus: {non_bus_users}")
  y -= 30

  p.drawString(50, y, "Students per Campus:")
  y -= 20
  for campus in campus_counts:
    p.drawString(70, y, f"{campus['campus_id__campus_name']}: {campus['count']}")
    y -= 20
    if y < 100:
      p.showPage()
      y = 800

  y -= 10
  p.drawString(50, y, "Detailed Student List:")
  y -= 30

  students = Student.objects.select_related('campus_id').all()
  for student in students:
    line = f"{student.studentNumber} - {student.name} {student.surname} - {student.studentEmail} - {student.campus_id.campus_name} - Bus: {'Yes' if Bus_Stats.viewed_at else 'No'}"
    p.drawString(50, y, line)
    y -= 20
    if y < 100:
      p.showPage()
      y = 800

  p.save()
  return response


def student_report_docx(request):
  doc = Document()
  doc.add_heading('Student Report', 0)

  # Metrics
  total_students = Student.objects.count()
  bus_users = Bus_Stats.objects.count()
  non_bus_users = Bus_Stats.objects.count()
  campus_counts = Student.objects.values('campus_id__campus_name').annotate(count=Count('studentNumber'))

  # Summary Section
  doc.add_paragraph(f'Total Students: {total_students}')
  doc.add_paragraph(f'Students Using Bus: {bus_users}')
  doc.add_paragraph(f'Students Not Using Bus: {non_bus_users}')
  doc.add_paragraph("Students per Campus:")
  for campus in campus_counts:
    doc.add_paragraph(f"• {campus['campus_id__campus_name']}: {campus['count']}", style='List Bullet')

  # Table of students
  doc.add_paragraph('Student Details:')
  table = doc.add_table(rows=1, cols=5)
  table.style = 'Table Grid'
  hdr_cells = table.rows[0].cells
  hdr_cells[0].text = 'Student No.'
  hdr_cells[1].text = 'Name'
  hdr_cells[2].text = 'Surname'
  hdr_cells[3].text = 'Email'
  hdr_cells[4].text = 'Campus'

  students = Student.objects.select_related('campus_id').all()
  for student in students:
    row_cells = table.add_row().cells
    row_cells[0].text = student.studentNumber
    row_cells[1].text = student.name
    row_cells[2].text = student.surname
    row_cells[3].text = student.studentEmail
    row_cells[4].text = student.campus_id.campus_name if student.campus_id else "N/A"

  # Response
  response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
  response['Content-Disposition'] = 'attachment; filename=student_report.docx'
  doc.save(response)
  return response


def student_report_csv(request):
  # Create HTTP response with CSV content type
  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = 'attachment; filename="student_report.csv"'

  writer = csv.writer(response)

  # Write report title
  writer.writerow(['Student Report'])

  # Write column headers
  writer.writerow(['Student Number', 'Name', 'Surname', 'Email', 'Campus', 'Uses Bus'])

  students = Student.objects.select_related('campus_id').all()

  for student in students:
    writer.writerow([
      student.studentNumber,
      student.name,
      student.surname,
      student.studentEmail,
      student.campus_id.campus_name if student.campus_id else 'N/A',
      'Yes' if Bus_Stats.objects.count() else 'No'
    ])

  return response


def full_report(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  students = Student.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)
  initials = request.session.get("initials")

  # Import your Student model accordingly

  return render(request, 'admin/full_report.html', {
    'admin_count': admins.count(),
    'event_count': events.count(),
    'bus_count': buses.count(),
    'schedule_count': schedules.count(),
    'recent_admins': recent_admins,
    'recent_events': recent_events,
    'recent_schedules': recent_schedules,
    'recent_actions': recent_actions,
    'total_students': students.count(),
    'initials': initials    # Added student count here
  })


def full_report_pdf(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = 'attachment; filename="full_report.pdf"'

  buffer = BytesIO()
  p = canvas.Canvas(buffer)
  y = 800

  def write_line(text, indent=100):
    nonlocal y
    if y < 50:
      p.showPage()
      y = 800
    p.drawString(indent, y, text)
    y -= 20

  write_line(f"System Report ({now().date()})")
  write_line(
    f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}")
  write_line("")

  write_line("Recent Admins:")
  for admin in recent_admins:
    write_line(f"{admin.name} {admin.surname} - {admin.email}", indent=120)

  write_line("Recent Events:")
  for event in recent_events:
    write_line(f"{event.description} at {event.location} on {event.date}", indent=120)

  write_line("Recent Schedules:")
  for sched in recent_schedules:
    write_line(f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})", indent=120)

  write_line("Admin Actions:")
  for action in recent_actions:
    write_line(f"{action.admin_id.name}: {action.action_type} at {action.datetime}", indent=120)

  p.showPage()
  p.save()

  pdf = buffer.getvalue()
  buffer.close()
  response.write(pdf)
  return response


def full_report_docx(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()
  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  doc = Document()
  doc.add_heading('System Report', 0)
  doc.add_paragraph(f"Date: {now().date()}")
  doc.add_paragraph(
    f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}")

  doc.add_heading("Recent Admins", level=1)
  for admin in recent_admins:
    doc.add_paragraph(f"{admin.name} {admin.surname} - {admin.email}")

  doc.add_heading("Recent Events", level=1)
  for event in recent_events:
    doc.add_paragraph(f"{event.description} at {event.location} on {event.date}")

  doc.add_heading("Recent Schedules", level=1)
  for sched in recent_schedules:
    doc.add_paragraph(f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})")

  doc.add_heading("Admin Actions", level=1)
  for action in recent_actions:
    doc.add_paragraph(f"{action.admin_id.name}: {action.action_type} at {action.datetime}")

  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)

  response = HttpResponse(buffer.getvalue(),
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
  response['Content-Disposition'] = 'attachment; filename="full_report.docx"'
  return response


def full_report_csv(request):
  one_week_ago = now() - timedelta(days=7)
  admins = Admin.objects.all()
  events = Event.objects.all()
  buses = Bus.objects.all()
  schedules = Bus_schedule.objects.all()

  recent_admins = admins.filter(created_at__gte=one_week_ago) if hasattr(Admin, 'created_at') else []
  recent_events = events.filter(created_at__gte=one_week_ago) if hasattr(Event, 'created_at') else []
  recent_schedules = schedules.filter(departure_time__gte=now().time())
  recent_actions = Admin_Action.objects.filter(datetime__gte=one_week_ago)

  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = 'attachment; filename="full_report.csv"'
  writer = csv.writer(response)

  writer.writerow(["Section", "Details"])
  writer.writerow(["Counts",
                   f"Admins: {admins.count()}, Events: {events.count()}, Buses: {buses.count()}, Schedules: {schedules.count()}"])

  writer.writerow(["Recent Admins"])
  for admin in recent_admins:
    writer.writerow(["", f"{admin.name} {admin.surname} - {admin.email}"])

  writer.writerow(["Recent Events"])
  for event in recent_events:
    writer.writerow(["", f"{event.description} at {event.location} on {event.date}"])

  writer.writerow(["Recent Schedules"])
  for sched in recent_schedules:
    writer.writerow(["", f"{sched.departure} to {sched.destination} ({sched.departure_time}-{sched.arrival_time})"])

  writer.writerow(["Admin Actions"])
  for action in recent_actions:
    writer.writerow(["", f"{action.admin_id.name}: {action.action_type} at {action.datetime}"])

  return response

def events_stats():

  #line graph values
  today = date.today()
  start_of_week = today - timedelta(days=today.weekday())  # Monday
  end_of_week = start_of_week + timedelta(days=6)  # Sunday

  week_dates = [start_of_week + timedelta(days=i) for i in range(7)]
  labels = [d.strftime("%Y-%m-%d") for d in week_dates]

  rsvp_counts = (
      RSVP.objects
      .filter(done_at__date__gte=start_of_week, done_at__date__lte=end_of_week)
      .annotate(day=TruncDate('done_at'))
      .values('day')
      .annotate(count=Count('id'))
      .order_by('day')
  )

  rsvp_data = {entry['day']: entry['count'] for entry in rsvp_counts}
  data = [rsvp_data.get(d, 0) for d in week_dates]

  total_attendees = RSVP.objects.filter(done_at__date__gte=start_of_week, done_at__date__lte=end_of_week).count()

  #bar graph values
  today = now().date()
  month_start = today.replace(day=1)
  next_month = (month_start.replace(day=28) + timedelta(days=4)).replace(day=1)

  events_data = Event.objects.filter(
    date__gte=month_start,
    date__lt=next_month
  )

  event_names = [event.title for event in events_data]
  rsvp_counts = [event.attendance_count for event in events_data]

  context = {
    'labels': labels,
    'data': data,
    'total_attendees':total_attendees,
    'event_names': json.dumps(event_names),  # Convert to JSON
    'rsvp_counts': json.dumps(rsvp_counts),
  }
  return context

def download_student_details_pdf(request):
  studentEmail=request.GET.get('studentEmail')
  student = Student.objects.get(studentEmail=studentEmail)

  response = HttpResponse(content_type='application/pdf')
  response['Content-Disposition'] = f'attachment; filename="student_{student.studentNumber}_details.pdf"'

  p = canvas.Canvas(response, pagesize=A4)
  width, height = A4
  y = height - 50

  # Title
  p.setFont("Helvetica-Bold", 16)
  p.drawString(50, y, "Student Details: ")
  y -= 15
  p.line(50, y, 550, y)
  y -= 5
  y-=15

  p.setFont("Helvetica", 12)
  p.drawString(50, y, f"Student Number: {student.studentNumber}")
  y -= 20
  p.drawString(50, y, f"Name: {student.name}")
  y -= 20
  p.drawString(50, y, f"Surname: {student.surname}")
  y -= 20
  p.drawString(50, y, f"Email: {student.studentEmail}")
  y -= 20
  p.drawString(50, y, f"Campus: {student.campus_id.campus_name if student.campus_id else 'N/A'}")
  y -= 20
  p.drawString(50, y,f"Last Login Time: {student.login_time.strftime('%Y-%m-%d %H:%M') if student.login_time else 'N/A'}")
  y -= 20

  y -= 15

  # Section 1: Event RSVPs
  p.setFont("Helvetica-Bold", 14)
  p.drawString(50, y, "Event RSVPs done by student")
  y -= 5
  p.line(50, y, 550, y)
  y -= 5
  y-=15
  p.setFont("Helvetica", 12)

  rsvps = RSVP.objects.filter(guest_studentnumber=student.studentNumber).order_by('done_at')

  if rsvps.exists():
        for rsvp in rsvps:
            if y < 100:
                p.showPage()
                y = height - 50
                p.setFont("Helvetica", 12)
            p.drawString(60, y, f"- {rsvp.done_at.strftime('%Y-%m-%d %H:%M')} : RSVP’d to {rsvp.event.title}")
            y -= 20
  else:
        p.drawString(60, y, "No event RSVPs.")
        y -= 20

  y -= 15

  # Section 2: Bus Schedule views
  p.setFont("Helvetica-Bold", 14)
  p.drawString(50, y, "Bus schedule views by student")
  y -= 5
  p.line(50, y, 550, y)
  y -= 5
  y -= 15
  p.setFont("Helvetica", 12)
  bus_views = Bus_Stats.objects.filter(student_id=student.studentNumber).order_by('viewed_at')

  if bus_views.exists():
      for view in bus_views:
          if y < 100:
              p.showPage()
              y = height - 50
              p.setFont("Helvetica", 12)
          p.drawString(60, y,f"- {view.viewed_at.strftime('%Y-%m-%d %H:%M')} : Viewed bus schedule of '{view.schedule_code}'")

          y -= 20
  else:
      p.drawString(60, y, "No bus schedule views.")
      y -= 20

  y -= 20

  p.save()
  return response

def download_student_details_csv(request):
  studentEmail = request.GET.get('studentEmail')
  student = Student.objects.get(studentEmail=studentEmail)

  response = HttpResponse(content_type='text/csv')
  response['Content-Disposition'] = f'attachment; filename="student_{student.studentNumber}_details.csv"'

  writer = csv.writer(response)
  # Header row
  writer.writerow(['Student Number', 'Name', 'Surname', 'Email', 'Campus', 'Last Login '])

  # Data row
  writer.writerow([
    student.studentNumber,
    student.name,
    student.surname,
    student.studentEmail,
    student.campus_id.campus_name if student.campus_id else "N/A",
    student.login_time.strftime('%Y-%m-%d %H:%M') if student.login_time else "N/A"
  ])

  writer.writerow(['Bus Schedule Views by student'])
  writer.writerow(['Timestamp', 'Schedule Code'])

  bus_views = Bus_Stats.objects.filter(student_id=student.studentNumber).order_by('viewed_at')
  if bus_views.exists():
      for view in bus_views:
          writer.writerow([view.viewed_at.strftime('%Y-%m-%d %H:%M'), view.schedule_code])
  else:
      writer.writerow(["No bus schedule."])

  writer.writerow([])  # empty row

  # Section 3: Event RSVPs
  writer.writerow(['Event RSVPs done by student'])
  writer.writerow(['Timestamp', 'Event Name'])

  rsvps = RSVP.objects.filter(guest_studentnumber=student.studentNumber).order_by('done_at')
  if rsvps.exists():
      for rsvp in rsvps:
          writer.writerow([rsvp.done_at.strftime('%Y-%m-%d %H:%M'), rsvp.event.title])
  else:
      writer.writerow(["No event RSVPs."])
  return response

def get_user_activity_data():
    today = date.today()
    start_of_week = today - timedelta(days=today.weekday())  # Monday
    week_dates = [start_of_week + timedelta(days=i) for i in range(7)]
    labels = [d.strftime("%Y-%m-%d") for d in week_dates]

    # Query for WAU
    wau_raw = (
        Student.objects
        .filter(login_time__date__gte=start_of_week, login_time__date__lte=today)
        .annotate(day=TruncDate('login_time'))
        .values('day')
        .annotate(count=Count('studentNumber'))
        .order_by('day')
    )

    # Create a date-to-count dictionary
    wau_map = {entry['day'].strftime('%Y-%m-%d'): entry['count'] for entry in wau_raw}

    # Ensure all 7 days are covered (fill missing days with 0)
    counts = [wau_map.get(day, 0) for day in labels]

    return {
        'wau_labels': json.dumps(labels),
        'wau_counts': json.dumps(counts)
    }

def admin_report(request):
    icon = request.GET.get('type')
    admin_id=request.session.get('admin_id')
    admin=Admin.objects.get(admin_id=admin_id)

    action_type='all'
    if icon=='bi bi-map':
        action_type='Map Actions'
    if icon=='bi bi-calendar':
        action_type='Event Actions'
    elif icon=='bi bi-bus-front':
        action_type='Bus Actions'

    if icon and icon!= 'all':
        actions = Admin_Action.objects.filter(icon=icon,admin_id=admin.admin_id)
    else:
        actions = Admin_Action.objects.all()

        # Create PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{action_type}_admin_report.pdf"'

    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4
    y = height

    # Title
    p.setFont("Helvetica", 12)
    p.drawString(50, y, f"Admin ID: {admin.admin_id}")
    y -= 20
    p.drawString(50, y, f"Name: {admin.name}")
    y -= 20
    p.drawString(50, y, f"Admin ID: {admin.admin_id}")
    y -= 20

    p.setFont("Helvetica-Bold", 16)
    p.drawCentredString(width / 2.0, height - 50, f"Admin Report - {action_type.upper()}")


    # Table headers
    y = height - 90
    p.setFont("Helvetica-Bold", 11)
    p.drawString(150, y, "Description")
    p.drawString(480, y, "Timestamp")

    y -= 20
    p.setFont("Helvetica", 10)

    for action in actions:
        if y < 50:
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 10)

        action_type = action.get_action_type_display() if hasattr(action,'get_action_type_display') else action.action_type
        timestamp = localtime(action.datetime).strftime('%Y-%m-%d %H:%M')

        p.drawString(150, y, action_type)
        p.drawString(480, y, timestamp)

        y -= 18

    p.showPage()
    p.save()

    return response
